import sys
import os
import pickle
import threading
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import base64
import uuid
import imgkit
from PIL import Image
from googleapiclient.errors import HttpError
from tkinter import filedialog, messagebox
from tkinter import ttk
import tkinter as tk
from tkinter import messagebox, ttk
import random
import string
import logging
from datetime import datetime
from docx import Document
import webbrowser

def generate_random_filename(mode):
    """Generate a random filename based on the selected mode."""
    length = random.randint(22, 25)
    
    if mode == 1:  # Mode 1: Digits only
        return ''.join(random.choices('0123456789', k=length))
    elif mode == 2:  # Mode 2: Letters only
        return ''.join(random.choices(string.ascii_uppercase, k=length))
    elif mode == 3:  # Mode 3: Digits and letters
        return ''.join(random.choices('0123456789' + string.ascii_uppercase, k=length))
    elif mode == 4:  # Mode 4: uuid
        return str(uuid.uuid4())
    else:
        raise ValueError("Invalid mode selected")

# Define separate directories for tokens and PDFs
TOKEN_DIR = 'tokens'
PDF_DIR = 'pdfs'

# Ensure the directories exist
os.makedirs(TOKEN_DIR, exist_ok=True)
os.makedirs(PDF_DIR, exist_ok=True)

SCOPES = [
    'https://www.googleapis.com/auth/gmail.send',
    'https://www.googleapis.com/auth/gmail.readonly',
    'https://www.googleapis.com/auth/gmail.modify',
    'https://www.googleapis.com/auth/gmail.compose',
    'https://www.googleapis.com/auth/gmail.labels',
    'https://www.googleapis.com/auth/gmail.metadata',
    'https://www.googleapis.com/auth/gmail.settings.basic',
    'https://www.googleapis.com/auth/gmail.settings.sharing'
]

# Resource path function to handle PyInstaller bundling
def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# Specify the path to wkhtmltoimage
wkhtmltoimage_exe = resource_path(os.path.join("resources", "wkhtmltoimage.exe"))  # Update for your OS

# Configure imgkit with the correct path to wkhtmltoimage
config = imgkit.config(wkhtmltoimage=wkhtmltoimage_exe)

# Configure logging
logging.basicConfig(
    filename=os.path.join(PDF_DIR, 'email_sender.log'),
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# Function to sanitize email for filename
def sanitize_email(email):
    return email.replace('@', '_at_').replace('.', '_dot_')

# Function to generate random DPI and image quality
def generate_random_pdf_options():
    dpi = random.choice([50, 60, 72, 80, 96, 120, 150, 175, 200, 225, 250, 275, 300])
    quality = random.choice([70, 72, 75, 78, 80, 82, 85, 87, 89, 90])
    zoom = random.uniform(0.75, 1.15)  # Zoom level between 0.75 and 1.15
    return dpi, quality, zoom

def html_to_jpg(html_content, jpg_filename):
    try:
        jpg_path = os.path.join(PDF_DIR, jpg_filename)
        dpi, quality, zoom = generate_random_pdf_options()

        options = {
            'format': 'jpg',
            'quality': str(quality),
            'zoom': str(zoom),
            'encoding': "UTF-8",
            'enable-local-file-access': '',
        }
        imgkit.from_string(html_content, jpg_path, options=options, config=config)
        logging.info(f"JPG generated: {jpg_path}")
        return jpg_path
    except Exception as e:
        logging.error(f"Error generating JPG: {e}")
        return None

import random
import logging
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

def jpg_to_doc(jpg_path, doc_path):
    try:
        # Create a new Document
        doc = Document()
        
        # Add a paragraph for the image
        paragraph = doc.add_paragraph()
        
        # Add the image to the paragraph with random width
        run = paragraph.add_run()
        random_width = random.choice([6.5 ,6.75 ,7, 7.5, 7.75, 8])  # Randomly select a width
        run.add_picture(jpg_path, width=Inches(random_width))
        
        # Center the paragraph
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Save the document
        doc.save(doc_path)
        logging.info(f"DOC generated: {doc_path} with image width: {random_width} inches")
        return doc_path
    except Exception as e:
        logging.error(f"Error converting JPG to DOC: {e}")
        return None


# Generate DOC from HTML content via JPG
def html_to_doc(html_content, doc_filename):
    try:
        # First convert HTML to JPG
        jpg_filename = f"{uuid.uuid4()}.jpg"
        jpg_path = html_to_jpg(html_content, jpg_filename)

        if not jpg_path or not os.path.exists(jpg_path):
            return None

        # Then convert JPG to DOC
        doc_path = os.path.join(PDF_DIR, doc_filename)
        doc_generated = jpg_to_doc(jpg_path, doc_path)

        if doc_generated:
            # Optionally delete the JPG after DOC is created
            os.remove(jpg_path)
            logging.info(f"Deleted intermediate JPG: {jpg_path}")
        
        return doc_path
    except Exception as e:
        logging.error(f"Error generating DOC: {e}")
        return None

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics

def html_to_pdf(html_content, pdf_filename):
    try:
        jpg_filename = f"{uuid.uuid4()}.jpg"
        jpg_path = html_to_jpg(html_content, jpg_filename)

        if not jpg_path or not os.path.exists(jpg_path):
            return None

        pdf_path = os.path.join(PDF_DIR, pdf_filename)

        # Create a PDF/A compliant document
        c = canvas.Canvas(pdf_path, pagesize=letter)
        c.setTitle("PDF/A Document")
        
        # Add the image to the PDF
        c.drawImage(jpg_path, 0, 0, width=8.5 * inch, height=11 * inch)  # Adjust size as needed

        # Add PDF/A compliance metadata
        c.setAuthor("Your Name")
        c.setSubject("PDF/A Compliance Example")
        c.setTitle("PDF/A Document Title")
        c.setCreator("Your Application Name")
        c.setKeywords("PDF/A, Example")

        # Finish the PDF
        c.showPage()
        c.save()

        logging.info(f"PDF generated: {pdf_path}")

        # Optionally delete the JPG after PDF is created
        os.remove(jpg_path)
        logging.info(f"Deleted intermediate JPG: {jpg_path}")

        return pdf_path
    except Exception as e:
        logging.error(f"Error generating PDF: {e}")
        return None

def create_message_with_html_doc_attachment(sender, to, subject, body, html_content, doc_filename):
    message = MIMEMultipart()
    message['to'] = to
    message['from'] = sender
    message['subject'] = subject

    # Attach the email body
    msg = MIMEText(body, 'html')
    message.attach(msg)

    message.add_header('X-Gmail-Message-ID', '1')             # Custom message ID for Gmail
    message.add_header('X-Priority', '1')               # High priority (1 = high, 5 = low)
    message.add_header('Importance', 'high')            # Generic importance header
    message.add_header('X-Important', '1')              # Custom header, some clients might respect it
    message.add_header('Priority', 'urgent')            # Outlook and older clients
    message.add_header('X-MSMail-Priority', 'High')     # Microsoft Outlook/Exchange
    message.add_header('X-MXPriority', '1')             
    # Adding more headers for Google
    message.add_header('X-Google-Original-Message-ID', '1')  # Custom header for Google
    message.add_header('X-Importance', 'high')                # Another importance header
    message.add_header('X-Gmail-Labels', 'Important')         # Label the email as important
    message.add_header('X-Notify', '1')                       # Custom header to trigger notifications
    message.add_header('X-Gmail-Thread-ID', '1')              # Thread ID for Gmail conversations
    message.add_header('X-Gmail-Message-Id', '1')             # Another custom message ID for Gmail

    if html_content.strip():  # Only proceed if html_content is not empty
        # Convert HTML to DOC via JPG
        doc_file = html_to_doc(html_content, doc_filename)

        if doc_file and os.path.exists(doc_file):
            with open(doc_file, 'rb') as f:
                mime_part = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
                mime_part.set_payload(f.read())
                encoders.encode_base64(mime_part)
                mime_part.add_header('Content-Disposition', f'attachment; filename="{os.path.basename(doc_file)}"')
                message.attach(mime_part)

    raw_message = base64.urlsafe_b64encode(message.as_bytes()).decode()
    return {'raw': raw_message}

# Function to create a message with PDF attachment
def create_message_with_pdf_attachment(sender, to, subject, body, pdf_path):
    message = MIMEMultipart()
    message['to'] = to
    message['from'] = sender
    message['subject'] = subject

    # Attach the email body
    msg = MIMEText(body, 'html')
    message.attach(msg)

    message.add_header('X-Gmail-Message-ID', '1')             # Custom message ID for Gmail
    message.add_header('X-Priority', '1')               # High priority (1 = high, 5 = low)
    message.add_header('Importance', 'high')            # Generic importance header
    message.add_header('X-Important', '1')              # Custom header, some clients might respect it
    message.add_header('Priority', 'urgent')            # Outlook and older clients
    message.add_header('X-MSMail-Priority', 'High')     # Microsoft Outlook/Exchange
    message.add_header('X-MXPriority', '1')             
    # Adding more headers for Google
    message.add_header('X-Google-Original-Message-ID', '1')  # Custom header for Google
    message.add_header('X-Importance', 'high')                # Another importance header
    message.add_header('X-Gmail-Labels', 'Important')         # Label the email as important
    message.add_header('X-Notify', '1')                       # Custom header to trigger notifications
    message.add_header('X-Gmail-Thread-ID', '1')              # Thread ID for Gmail conversations
    message.add_header('X-Gmail-Message-Id', '1')             # Another custom message ID for Gmail
    
    with open(pdf_path, 'rb') as f:
        mime_part = MIMEBase('application', 'pdf')
        mime_part.set_payload(f.read())
        encoders.encode_base64(mime_part)
        mime_part.add_header('Content-Disposition', f'attachment; filename="{os.path.basename(pdf_path)}"')
        message.attach(mime_part)

    raw_message = base64.urlsafe_b64encode(message.as_bytes()).decode()
    return {'raw' : raw_message}

# Placeholder Functions for Random Data
def generate_random_number(length=8):
    return ''.join(random.choices(string.digits, k=length))

def generate_random_snumber(length=4):
    return ''.join(random.choices(string.digits, k=length))

def generate_random_abc(length=5):
    return ''.join(random.choices(string.ascii_uppercase, k=length))

def generate_random_invoice(length=8):
    return ''.join(random.choices(string.ascii_uppercase + string.digits, k=length))

def generate_random_symbol():
    return random.choice(string.punctuation)

def generate_random_paragraph():
    paragraphs = [
        "If you feel like this transaction was not authorized by you or your account has been accessed fraudulently, then please call our toll-free number.",
    ]
    return random.choice(paragraphs)

def generate_random_paragrapha():
    paragraph = [
        "Bill",
    ]
    return random.choice(paragraph)

# Function to generate a serial number in the format GTIG-4RTIO-TGHU6-HGDA
def generate_random_serial_number():
    parts = []
    for _ in range(4):  # Create 4 parts
        part = ''.join(random.choices(string.ascii_uppercase, k=5))  # 5 uppercase letters
        if _ == 1:  # Second part should be a mix of letters and digits
            part = ''.join(random.choices(string.ascii_uppercase + string.digits, k=5))
        parts.append(part)
    return '-'.join(parts)

# Placeholder function now accepts recipient email to replace #EMAIL#
def replace_placeholders(text, recipient_email=None, tfn=None):
    text = text.replace("#NUMBER#", generate_random_number(random.randint(8, 16)))
    text = text.replace("#SNUMBER#", generate_random_snumber(random.randint(4, 8)))
    text = text.replace("#ABC#", generate_random_abc(random.randint(5, 8)))
    text = text.replace("#INVOICE#", generate_random_invoice(random.randint(8, 10)))
    text = text.replace("#SYMBOL#", generate_random_symbol())
    text = text.replace("#CONTENT#", generate_random_paragraph())
    text = text.replace("#HEADER#", generate_random_paragrapha())
    text = text.replace("#SIRIALNO#", generate_random_serial_number())  # Replace #SIRIALNO#

    if recipient_email:
        text = text.replace("#EMAIL#", recipient_email)
    if tfn:
        text = text.replace("#TFN#", tfn)
    
    # Add current date in YYYY-MM-DD format
    current_date = datetime.now().strftime('%Y-%m-%d')
    text = text.replace("#DATE#", current_date)

    return text

# Authentication Function (now threaded)
def authenticate_gmail(credentials_file, token_filename, callback):
    def authenticate():
        creds = None
        token_path = os.path.join(TOKEN_DIR, token_filename)  # Dynamic token path

        if os.path.exists(token_path):
            with open(token_path, 'rb') as token:
                creds = pickle.load(token)

        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                creds.refresh(Request())
            else:
                flow = InstalledAppFlow.from_client_secrets_file(credentials_file, SCOPES)
                creds = flow.run_local_server(port=0)
            with open(token_path, 'wb') as token:
                pickle.dump(creds, token)

        service = build('gmail', 'v1', credentials=creds)
        callback(service)

    threading.Thread(target=authenticate).start()

# Send Email Function
def send_email_with_doc(service, sender, to, subject, body, html_content, doc_filename, callback):
    try:
        message = create_message_with_html_doc_attachment(sender, to, subject, body, html_content, doc_filename)
        send_message = service.users().messages().send(userId='me', body=message).execute()
        print(f"Email sent to {to}")
        
        # Delete the DOC after sending
        doc_path = os.path.join(PDF_DIR, doc_filename)
        if os.path.exists(doc_path):
            os.remove(doc_path)
            print(f"Deleted DOC: {doc_path}")
        
        callback(to)
    except HttpError as error:
        print(f"An error occurred: {error}")
        callback(None)

def send_email_with_pdf(service, sender, recipient, subject, body, pdf_path, callback):
    try:
        message = create_message_with_pdf_attachment(sender, recipient, subject, body, pdf_path)
        send_message = service.users().messages().send(userId='me', body=message).execute()
        print(f"Email sent to {recipient}")

        # Delete the PDF after sending
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
            print(f"Deleted PDF: {pdf_path}")

        callback(recipient)
    except HttpError as error:
        print(f"An error occurred: {error}")
        callback(None)


# Modify the send_emails function to work with TaskWindow instances
def send_emails(credentials_file, sender_name, sender_email, subject, email_body, html_content, recipients, on_emails_sent, task_window):
    # Sanitize sender_email for filename
    sanitized_email = sanitize_email(sender_email)
    token_filename = f"token_{sanitized_email}.pickle"

    def send_all_emails(service):
        sender = f"{sender_name} <{sender_email}>"
        task_window.send_emails_threaded(service, sender, subject, email_body, html_content, recipients, on_emails_sent)

    authenticate_gmail(credentials_file, token_filename, send_all_emails)

# Class representing each Task as a Tab
import tkinter as tk
from tkinter import ttk
from tkinter import font

class TaskTab:
    task_count = 0
    max_tasks = 2

    def __init__(self, notebook, task_id, close_callback):
        self.notebook = notebook
        self.task_id = task_id
        self.close_callback = close_callback
        self.sent_count = 0
        self.total_recipients = 0
        self.style = ttk.Style()
        self.style.configure(f"TaskTab.TFrame", background="#cee6e6")
        self.frame = ttk.Frame(notebook, style="TaskTab.TFrame")
        self.notebook.add(self.frame, text=f"â™¨ï¸Task {self.task_id}  \n Sent: {self.sent_count}")
        self.create_widgets()
        self.stop_flag = False
        self.status = "In Progress"  # Track task status
        self.update_tab_title()  # Update title on initialization
        

    def update_tab_title(self):
        """Update the tab title to reflect the current sent count and status."""
        if self.status == "Completed":
            self.notebook.tab(self.frame, text=f"â™¨ï¸Task {self.task_id}  \nâœ…  Sent: {self.sent_count}")
        elif self.status == "Error":
            self.notebook.tab(self.frame, text=f"â™¨ï¸Task {self.task_id}  \nâŒ  Sent: {self.sent_count}")
        else:  # In Progress
            self.notebook.tab(self.frame, text=f"â™¨ï¸Task {self.task_id}  \nâ³  Sent: {self.sent_count}")

    def increment_sent_count(self):
        self.sent_count += 1
        self.update_tab_title()


    def create_widgets(self):
         # Define a custom font
        custom_font = font.Font(family="Arial", size=12)
        # Configure grid weights for responsive resizing
        self.frame.columnconfigure(1, weight=1)

        # Sender Email
        tk.Label(self.frame, text="Sender Email:", bg="#cee6e6").grid(row=0, column=0, padx=1, pady=1, sticky='e')
        self.sender_email_entry = tk.Entry(self.frame, width=5, bg='lightgray', font=custom_font)
        self.sender_email_entry.grid(row=0, column=1, padx=1, pady=1, sticky='we', ipadx=1)
        tk.Button(self.frame, text="Copy", command=lambda: self.copy_text(self.sender_email_entry), bg='lightgreen').grid(row=0, column=2, padx=1, pady=1)

        # Sender Name
        tk.Label(self.frame, text="Sender Name:", bg="#cee6e6").grid(row=1, column=0, padx=1, pady=1, sticky='e')
        self.sender_name_entry = tk.Entry(self.frame, width=5, bg='lightgray', font=custom_font)
        self.sender_name_entry.grid(row=1, column=1, padx=1, pady=1, sticky='we', ipadx=1)
        tk.Button(self.frame, text="Paste", command=lambda: self.paste_to_entry(self.sender_name_entry), bg='lightgreen').grid(row=1, column=2, padx=1, pady=1)
        
        # Subject
        tk.Label(self.frame, text="Subject:", bg="#cee6e6").grid(row=2, column=0, padx=1, pady=1, sticky='e')
        self.subject_entry = tk.Entry(self.frame, width=5, bg='lightgray', font=custom_font)
        self.subject_entry.grid(row=2, column=1, padx=1, pady=1, sticky='we', ipadx=1)
        tk.Button(self.frame, text="Paste", command=lambda: self.paste_to_entry(self.subject_entry), bg='lightgreen').grid(row=2, column=2, padx=1, pady=1)

        # Email Body
        tk.Label(self.frame, text="Email Body:", bg="#cee6e6").grid(row=3, column=0, padx=1, pady=1, sticky='ne')
        self.email_body_entry = tk.Text(self.frame, width=5, height=2, bg='lightgray')
        self.email_body_entry.grid(row=3, column=1, padx=1, pady=1, sticky='we', ipadx=1)
        tk.Button(self.frame, text="Paste", command=lambda: self.paste_text(self.email_body_entry), bg='lightgreen').grid(row=3, column=2, padx=1, pady=1)

        # HTML Content
        tk.Label(self.frame, text="HTML Con:", bg="#cee6e6").grid(row=4, column=0, padx=1, pady=1, sticky='ne')
        self.html_content_entry = tk.Text(self.frame, width=5, height=4, bg='lightgray')
        self.html_content_entry.grid(row=4, column=1, padx=1, pady=1, sticky='we', ipadx=1)
        tk.Button(self.frame, text="Paste", command=lambda: self.paste_text(self.html_content_entry), bg='lightgreen').grid(row=4, column=2, padx=1, pady=1)
        tk.Button(self.frame, text="Preview HTML", command=self.view_html, bg='#f7e887').grid(row=10, column=1, padx=1, pady=1, sticky='e')
        # Recipient List
        tk.Label(self.frame, text="Recipient List:", bg="#cee6e6").grid(row=5, column=0, padx=1, pady=1, sticky='ne')
        self.recipient_list_entry = tk.Text(self.frame, width=5, height=4, bg='lightgray')
        self.recipient_list_entry.grid(row=5, column=1, rowspan=2, padx=1, pady=1, sticky='we', ipadx=1)
        tk.Button(self.frame, text="Paste", command=lambda: self.paste_text(self.recipient_list_entry), bg='lightgreen').grid(row=5, column=2, padx=1, pady=1)
        tk.Button(self.frame, text="Get", command=self.insert_email, bg='lightgreen').grid(row=6, column=2, padx=1, pady=1)

        # OAuth2 Credentials
        tk.Label(self.frame, text="API:", bg="#cee6e6").grid(row=7, column=0, padx=1, pady=1, sticky='e')
        self.credentials_path = tk.StringVar()
        self.credentials_entry = tk.Entry(self.frame, textvariable=self.credentials_path, width=5, state=tk.DISABLED, bg='lightgray')
        self.credentials_entry.grid(row=7, column=1, padx=1, pady=1, sticky='we', ipadx=1)
        tk.Button(self.frame, text="Load", command=self.upload_credentials, bg='lightyellow').grid(row=7, column=2, padx=1, pady=1)
        
        # TFN Input
        tk.Label(self.frame, text="TFN:", bg="#cee6e6").grid(row=8, column=0, padx=1, pady=1, sticky='e')
        self.tfn_entry = tk.Entry(self.frame, width=5, bg='lightgray', font=custom_font)
        self.tfn_entry.grid(row=8, column=1, padx=1, pady=1, sticky='we', ipadx=1)

        self.conversion_type = tk.StringVar(value="html_to_doc")  # Default conversion type

        # Conversion Type
        tk.Radiobutton(self.frame, text="To DOC", variable=self.conversion_type, value="html_to_doc", bg="#cee6e6").grid(row=14, column=0, padx=1, pady=1, sticky='w')
        tk.Radiobutton(self.frame, text="To PDF", variable=self.conversion_type, value="html_to_pdf", bg="#cee6e6").grid(row=15, column=0, padx=1, pady=1, sticky='w')

        # Progress Bar
        self.progress_bar = ttk.Progressbar(self.frame, orient=tk.HORIZONTAL, length=400, mode='determinate')
        self.progress_bar.grid(row=9, column=1, pady=1, sticky='we', ipadx=1)

        # Email Count Label
        self.email_count_label = tk.Label(self.frame, text="ðŸ“§Emails Sent: 0/0 (0%)", bg="#cee6e6")
        self.email_count_label.grid(row=10, column=1, pady=0, sticky='w')

        # Current Recipient Label
        self.current_recipient_label = tk.Label(self.frame, text="", bg="#cee6e6")
        self.current_recipient_label.grid(row=11, column=1, padx=1, pady=1, sticky='we')

        # Send and Stop Buttons
        tk.Button(self.frame, text="ðŸ“¤ Send", command=self.send_email_process, bg='green', fg='white').grid(row=12, column=1, pady=1, sticky='w')
        tk.Button(self.frame, text="â¹ Stop", command=self.stop_sending, bg='red', fg='white').grid(row=12, column=1, pady=1, sticky='e')

        # Close Button
        tk.Button(self.frame, text="âŒ Close", bg='pink', fg='black', command=lambda: self.close_callback(self)).grid(row=13, column=0, padx=1, pady=1)
        # Status Text
        self.status_text = tk.Text(self.frame, height=7, width=5, state=tk.DISABLED, bg='lightgray')
        self.status_text.grid(row=13, column=1, rowspan=4, padx=1, pady=(1, 5), sticky='we', ipadx=1)  # Adjusted row to 10

        # GMass Button
        tk.Button(self.frame, text="Get GMass", command=self.send_gmass_emails, bg='blue', fg='white').grid(row=12, column=1, padx=1, pady=1)
        # Filename Mode
        self.filename_mode = tk.IntVar(value=1)  # Default to Mode 1
        tk.Radiobutton(self.frame, text="M-1", variable=self.filename_mode, value=1, bg="#cee6e6").grid(row=13, column=2, padx=1, pady=1, sticky='w')
        tk.Radiobutton(self.frame, text="M-2", variable=self.filename_mode, value=2, bg="#cee6e6").grid(row=14, column=2, padx=1, pady=1, sticky='w')
        tk.Radiobutton(self.frame, text="M-3", variable=self.filename_mode, value=3, bg="#cee6e6").grid(row=15, column=2, padx=1, pady=1, sticky='w')
        tk.Radiobutton(self.frame, text="M-4", variable=self.filename_mode, value=4, bg="#cee6e6").grid(row=16, column=2, padx=1, pady=1, sticky='w')
    def insert_email(self):
            self.recipient_list_entry.delete(1.0, tk.END)  # Clear existing text
            self.recipient_list_entry.insert(tk.END, "littlesayn@gmail.com")  # Insert the email
            self.status_text.grid_forget()
            # Clear the status text before sending emails
            self.status_text.configure(state=tk.NORMAL)
            self.status_text.delete("1.0", tk.END)
            self.status_text.configure(state=tk.DISABLED)

            # Re-add the status text to the grid
            self.status_text.grid(row=13, column=1, rowspan=4, padx=1, pady=(1, 5), sticky='we', ipadx=1)
            self.progress_bar['value'] = 0  # Reset progress bar
            self.progress_bar.update()
            self.email_count_label.config(text="ðŸ“¨ Emails Sent: 0/0 (0%)")  # Reset email count label

    def view_html(self):
        html_content = self.html_content_entry.get("1.0", tk.END).strip()
        if not html_content:
            messagebox.showwarning("Warning", "No HTML content to view!")
            return

        # Create a temporary HTML file
        html_filename = os.path.join(PDF_DIR, f"temp_view_{uuid.uuid4()}.html")
        with open(html_filename, 'w', encoding='utf-8') as f:
            f.write(html_content)

        # Open the HTML file in the default web browser
        webbrowser.open(f'{html_filename}')

    def upload_credentials(self):
        file_path = filedialog.askopenfilename(filetypes=[("JSON files", "*.json")])
        if file_path:
            self.credentials_path.set(file_path)
            file_name = os.path.basename(file_path)
            sender_email = os.path.splitext(file_name)[0]
            self.sender_email_entry.delete(0, tk.END)
            self.sender_email_entry.insert(0, sender_email)

    def update_status(self, message):
        self.status_text.configure(state=tk.NORMAL)
        self.status_text.insert(tk.END, message + "\n")
        self.status_text.yview(tk.END)
        self.status_text.configure(state=tk.DISABLED)
        
    def paste_to_entry(self, entry):
        entry.delete(0, tk.END)  # Clear the entry
        entry.insert(0, self.frame.clipboard_get())  # Paste from clipboard

    def paste_text(self, text_widget):
        try:
            clipboard_content = self.frame.clipboard_get()
            text_widget.delete("1.0", tk.END)
            text_widget.insert("1.0", clipboard_content)
            self.status_text.grid_forget()
            self.status_text.configure(state=tk.NORMAL)
            self.status_text.delete("1.0", tk.END)
            self.status_text.configure(state=tk.DISABLED)
            self.progress_bar['value'] = 0  # Reset progress bar
            self.progress_bar.update()
            self.email_count_label.config(text="ðŸ“¨ Emails Sent: 0/0 (0%)")  # Reset email count label

            # Re-add the status text to the grid
            self.status_text.grid(row=13, column=1, rowspan=4, padx=1, pady=(1, 5), sticky='we', ipadx=1)
        except tk.TclError:
            messagebox.showwarning("Paste Error", "Clipboard is empty or contains non-text data.")

    def copy_text(self, entry_widget):
        self.frame.clipboard_clear()
        self.frame.clipboard_append(entry_widget.get())

    def stop_sending(self):
        self.stop_flag = True
        self.update_status("â¹ Email sending stopped.")

    def send_gmass_emails(self):
            self.recipient_list_entry.delete(1.0, tk.END)  # Clear existing text
            self.recipient_list_entry.insert(tk.END, "ajaygoel999@gmail.com\ntest@chromecompete.com\ntest@ajaygoel.org\nme@dropboxslideshow.com\ntest@wordzen.com\nrajgoel8477@gmail.com\nrajanderson8477@gmail.com\nrajwilson8477@gmail.com\nbriansmith8477@gmail.com\noliviasmith8477@gmail.com\nashsmith8477@gmail.com\nshellysmith8477@gmail.com\najay@madsciencekidz.com\najay2@ctopowered.com\najay@arena.tec.br\najay@daustin.co")  # Insert the email
            self.status_text.grid_forget()
            # Clear the status text before sending emails
            self.status_text.configure(state=tk.NORMAL)
            self.status_text.delete("1.0", tk.END)
            self.status_text.configure(state=tk.DISABLED)

            # Re-add the status text to the grid
            self.status_text.grid(row=13, column=1, rowspan=4, padx=1, pady=(1, 5), sticky='we', ipadx=1)
            self.progress_bar['value'] = 0  # Reset progress bar
            self.progress_bar.update()
            self.email_count_label.config(text="ðŸ“¨ Emails Sent: 0/0 (0%)")  # Reset email count label
    def send_email_process(self):
        # Reset sent count and update tab title
        self.sent_count = 0
        self.status = "In Progress"  # Set status to In Progress
        self.update_tab_title()  # Update the tab title to reflect the In Progress status

        sender_email = self.sender_email_entry.get().strip()
        sender_name = self.sender_name_entry.get().strip()
        subject = self.subject_entry.get().strip()
        email_body = self.email_body_entry.get("1.0", tk.END).strip()
        html_content = self.html_content_entry.get("1.0", tk.END).strip()
        credentials_file = self.credentials_path.get().strip()
        tfn = self.tfn_entry.get().strip()  # Get the TFN value
        recipients = [email.strip() for email in self.recipient_list_entry.get("1.0", tk.END).strip().splitlines() if email.strip()]

        self.status_text.grid_forget()
        self.status_text.configure(state=tk.NORMAL)
        self.status_text.delete("1.0", tk.END)
        self.status_text.configure(state=tk.DISABLED)
        self.progress_bar['value'] = 0  # Reset progress bar
        self.progress_bar.update()
        self.email_count_label.config(text="ðŸ“¨ Emails Sent: 0/0 (0%)")  # Reset email count label

        # Re-add the status text to the grid
        self.status_text.grid(row=13, column=1, rowspan=4, padx=1, pady=(1, 5), sticky='we', ipadx=1)

        # Check if all fields are filled
        if not credentials_file or not sender_email or not recipients:
            self.update_status("âš ï¸ Input Error: All fields must be filled except Subject and Body!")
            return

        # Set default subject and body if they are empty
        if not subject:
            subject = ""  # Default subject if empty
        if not email_body:
            email_body = ""  # Default body if empty

        self.total_recipients = len(recipients)
        self.progress_bar['value'] = 0  # Reset progress bar
        self.progress_bar.update()
        self.email_count_label.config(text=f"ðŸ“¨ Emails Sent: 0/{self.total_recipients} (0%)")  # Reset email count label
        self.stop_flag = False  # Ensure the stop flag is set to False

        def on_emails_sent(recipient):
            if recipient:
                self.update_status(f"âœ”ï¸ {recipient}")
                self.current_recipient_label.config(text=f"Sent to: {recipient} âœ”ï¸", fg="green")  # Show recipient with checkmark in green
            else:
                self.update_status("âŒ Failed to send email.")
                self.status = "Error"  # Set status to Error if any email fails

            self.increment_sent_count()

            progress_percentage = (self.sent_count / self.total_recipients) * 100
            self.progress_bar['value'] = progress_percentage
            self.progress_bar.update()
            self.email_count_label.config(text=f"ðŸ“¨ Emails Sent: {self.sent_count}/{self.total_recipients} ({int(progress_percentage)}%)")

            if self.sent_count >= self.total_recipients:
                self.status = "Completed"  # Set status to Completed if all emails sent
                self.update_tab_title()
                self.update_status("âœ… All emails processed.")

        # Start sending emails based on the selected conversion type
        send_emails(credentials_file, sender_name, sender_email, subject, email_body, html_content, recipients, on_emails_sent, self)

    def send_emails_threaded(self, service, sender, subject, email_body, html_content, recipients, callback):
        def send():
            for recipient in recipients:
                if self.stop_flag:
                    self.update_status("ðŸ“› Email sending stopped.")
                    break

                # Generate unique values for placeholders for each recipient
                personalized_body = replace_placeholders(email_body, recipient_email=recipient)
                personalized_html = replace_placeholders(html_content, recipient_email=recipient)
                personalized_subject = replace_placeholders(subject, recipient_email=recipient)

                # Determine the conversion type and generate the corresponding document
                if self.conversion_type.get() == "html_to_doc":
                    # Generate a unique DOC filename for each recipient based on selected mode
                    doc_filename = f"{generate_random_filename(self.filename_mode.get())}.docx"
                    send_email_with_doc(service, sender, recipient, personalized_subject, personalized_body, personalized_html, doc_filename, callback)
                elif self.conversion_type.get() == "html_to_pdf":
                    # Generate a unique PDF filename for each recipient based on selected mode
                    pdf_filename = f"{generate_random_filename(self.filename_mode.get())}.pdf"
                    pdf_path = html_to_pdf(personalized_html, pdf_filename)
                    if pdf_path:
                        send_email_with_pdf(service, sender, recipient, personalized_subject, personalized_body, pdf_path, callback)

        threading.Thread(target=send).start()

class LoginWindow:
    logged_in_user = None  # Class variable to track the logged-in user

    def __init__(self, master):
        self.master = master
        self.master.title("Login")
        self.master.geometry("300x200")

        self.label_id = tk.Label(master, text="ID:")
        self.label_id.pack(pady=5)

        self.entry_id = tk.Entry(master)
        self.entry_id.pack(pady=5)

        self.label_password = tk.Label(master, text="Password:")
        self.label_password.pack(pady=5)

        self.entry_password = tk.Entry(master, show="*")
        self.entry_password.pack(pady=5)

        self.login_button = tk.Button(master, text="Login", command=self.login)
        self.login_button.pack(pady=20)

        # Dictionary to hold valid user credentials
        self.valid_credentials = {
            "adhip1": "omg231",
            "user2": "password2",
            "user3": "password3",
            # Add more users as needed
        }

    def login(self):
        user_id = self.entry_id.get()
        password = self.entry_password.get()

        # Check if a user is already logged in
        if LoginWindow.logged_in_user is not None:
            messagebox.showwarning("Already Logged In", f"User  '{LoginWindow.logged_in_user}' is already logged in.")
            return

        # Check if the entered credentials are valid
        if user_id in self.valid_credentials and self.valid_credentials[user_id] == password:
            LoginWindow.logged_in_user = user_id  # Set the logged-in user
            self.master.destroy()  # Close the login window
            self.open_main_app()   # Open the main application
        else:
            messagebox.showerror("Login Failed", "Invalid ID or Password")

    def open_main_app(self):
        root = tk.Tk()
        root["bg"] = "#cee6e6"
        app = EmailSenderApp(root)
        root.mainloop()

class EmailSenderApp:
    def __init__(self, master):
        self.master = master
        self.master.title("ðŸ«• Pressure Cooker (PC_V_2.3)")
    
        # Set the window size (width x height)
        self.master.geometry("410x605+700+20")

        self.tasks = []
        self.create_main_widgets()
        # Initialize with one task
        self.add_task()

    def create_main_widgets(self):
        # Create a Frame for the button at the top
        button_frame = tk.Frame(self.master)
        button_frame.pack(side=tk.TOP, anchor='nw', padx=0, pady=0)  # Pack on the top side

        # Add Task Button
        add_task_button = tk.Button(button_frame, text="ð“Œ‰â—¯ð“‡‹ Add Task", command=self.add_task, bg='#9b9df3', fg='black')
        add_task_button.pack(side=tk.LEFT, padx=1, pady=1)  # Pack on the left side

        # Add Define All Tags Button on the opposite side
        define_tags_button = tk.Button(button_frame, text="Define All Tags & Descriptions", command=self.open_define_tags_window, bg='orange', fg='black')
        define_tags_button.pack(side=tk.RIGHT, padx=1, pady=1)  # Pack on the right side

        # Create a Notebook for tabs
        self.notebook = ttk.Notebook(self.master)
        self.notebook.pack(side=tk.TOP, anchor='nw', expand=1, fill='both')  # Pack below the button frame

    def open_define_tags_window(self):
        # Create a new window for defining tags
        tags_window = tk.Toplevel(self.master)
        tags_window.title("Define All Tags")
        tags_window.geometry("600x300")  # Set the size of the window

        # Create a Text widget to display the tags
        tags_text = tk.Text(tags_window, wrap=tk.WORD)
        tags_text.pack(expand=True, fill='both', padx=10, pady=10)

        # Define tags and their descriptions in a table format
        tags_info = (
            "{:<12} : {}\n"
            "{:<12} : {}\n"
            "{:<12} : {}\n"
            "{:<12} : {}\n"
            "{:<12} : {}\n"
            "{:<12} : {}\n"
            "{:<12} : {}\n"
            "{:<12} : {}\n"
            "{:<12} : {}\n"
            "{:<12} : {}\n"
            "{:<12} : {}\n"
            "{:<12} : {}\n"
            "{:<12} : {}\n"
            "{:<12} : {}\n"
            "{:<12} : {}\n"
        ).format(
            "#NUMBER#", "A random number, usually for identification.",
            "#SNUMBER#", "A serial number, often shorter than #NUMBER#.",
            "#ABC#", "A random string of uppercase letters.",
            "#INVOICE#", "A placeholder for invoice numbers.",
            "#SYMBOL#", "A random punctuation symbol.",
            "#HEADER#", "A random Header from predefined options.",
            "#CONTENT#", "Another random Content from predefined options.",
            "#SIRIALNO#", "A randomly generated serial number in a specific format.",
            "#EMAIL#", "The recipient's email address.",
            "#DATE#", "The current date in YYYY-MM-DD format.",
            "#TFN#", "The TFN number which you input.",
            "M-1", "Digits only for filename.",
            "M-2", "Uppercase letters only for filename.",
            "M-3", "Digits and uppercase letters for filename.",
            "M-4", "uudid for filename."
        )

        # Insert the tags info into the Text widget
        tags_text.insert(tk.END, tags_info)
        tags_text.configure(state=tk.DISABLED)  # Make the Text widget read-only

        # Add a close button
        close_button = tk.Button(tags_window, text="Close", command=tags_window.destroy)
        close_button.pack(pady=5)

    def add_task(self):
        if len(self.tasks) >= TaskTab.max_tasks:
            messagebox.showwarning("Task Limit Reached", f"You can only have up to {TaskTab.max_tasks} tasks.")
            return
        TaskTab.task_count += 1
        task = TaskTab(self.notebook, TaskTab.task_count, self.close_task)
        self.tasks.append(task)

    def close_task(self, task):
        self.notebook.forget(task.frame)
        self.tasks.remove(task)
        TaskTab.task_count -= 1

# Initialize and run the login application
if __name__ == "__main__":
    login_root = tk.Tk()
    login_app = LoginWindow(login_root)
    login_root.mainloop()

